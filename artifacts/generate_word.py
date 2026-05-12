import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Khởi tạo tài liệu
doc = Document()

# Tiêu đề chính
title = doc.add_heading('Báo Cáo Sự Cố Hệ Thống (Incident Reports)', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Thông tin chung
doc.add_paragraph('Vai trò: QA / SRE Engineer', style='List Bullet')
doc.add_paragraph('Dự án: Quản Lý Chi Tiêu (Expense Tracker)', style='List Bullet')

doc.add_paragraph('_' * 50)

# --- Incident 01 ---
doc.add_heading('Incident 01: Lỗi Crash Backend khi khởi động (ClassNotFoundException)', level=1)

doc.add_heading('1. Hiện tượng (Phenomenon)', level=2)
doc.add_paragraph('Mô tả: Khi thực hiện lệnh `mvn spring-boot:run` để khởi động server backend, tiến trình build thất bại và server bị crash ngay lập tức với mã lỗi `exit code: 1`.', style='List Bullet')
doc.add_paragraph('Log hệ thống ghi nhận:\nCaused by: java.lang.ClassNotFoundException: com.example.demo.dto.RegisterRequest', style='List Bullet')
doc.add_paragraph('Mức độ ảnh hưởng: Nghiêm trọng (Critical) - Toàn bộ dịch vụ backend không thể khởi động, chặn toàn bộ quá trình test và phát triển của team.', style='List Bullet')

doc.add_heading('2. Nguyên nhân (Root Cause)', level=2)
doc.add_paragraph('Mặc dù file mã nguồn `RegisterRequest.java` vẫn tồn tại đúng vị trí, nhưng thư mục `target` (chứa các file `.class` đã được biên dịch) bị lỗi bộ nhớ đệm (stale cache) hoặc out-of-sync.', style='List Bullet')
doc.add_paragraph('Trình biên dịch của Maven đã bỏ qua việc build lại class này, dẫn đến việc Spring Boot không thể tìm thấy nó trên Java Classpath khi chạy.', style='List Bullet')
doc.add_paragraph('Cách xử lý đã áp dụng: Xóa hoàn toàn thư mục build cũ và ép biên dịch lại từ đầu bằng lệnh `mvn clean compile`.', style='List Bullet')

doc.add_paragraph('_' * 50)

# --- Incident 02 ---
doc.add_heading('Incident 02: Xung đột tài nguyên mạng (Port 8081 already in use)', level=1)

doc.add_heading('1. Hiện tượng (Phenomenon)', level=2)
doc.add_paragraph('Mô tả: Sau khi khắc phục lỗi biên dịch (Incident 01), tiến trình khởi động Spring Boot tiếp tục thất bại ở giai đoạn khởi tạo Tomcat Web Server.', style='List Bullet')
doc.add_paragraph('Log hệ thống ghi nhận:\nAPPLICATION FAILED TO START\nDescription: Web server failed to start. Port 8081 was already in use.', style='List Bullet')
doc.add_paragraph('Mức độ ảnh hưởng: Cao (High) - Backend không thể binding vào cổng mạng được chỉ định, từ chối phục vụ các request từ Frontend.', style='List Bullet')

doc.add_heading('2. Nguyên nhân (Root Cause)', level=2)
doc.add_paragraph('Một tiến trình (Process) ngầm cũ của ứng dụng Spring Boot (PID: 13392) đã bị crash trước đó hoặc bị tắt không đúng cách (không giải phóng tài nguyên mạng).', style='List Bullet')
doc.add_paragraph('Do tiến trình zombie này vẫn đang "lắng nghe" (LISTENING) trên cổng `8081`, hệ điều hành từ chối cấp phát lại cổng này cho tiến trình mới, gây ra lỗi `BindException`.', style='List Bullet')
doc.add_paragraph('Cách xử lý đã áp dụng: Sử dụng `netstat -ano` để tìm PID đang chiếm dụng và dùng `taskkill /F /PID` để cưỡng chế dừng tiến trình zombie, giải phóng cổng mạng.', style='List Bullet')

doc.add_paragraph('_' * 50)

# --- Incident 03 ---
doc.add_heading('Incident 03: Vỡ Layout UI (Overflow) trên thiết bị di động', level=1)

doc.add_heading('1. Hiện tượng (Phenomenon)', level=2)
doc.add_paragraph('Mô tả: Trên trang Cài đặt (Settings), khi người dùng truy cập bằng thiết bị di động hoặc thu nhỏ cửa sổ trình duyệt, giao diện bị vỡ nghiêm trọng.', style='List Bullet')
doc.add_paragraph('Cụ thể: Khu vực "Vô hiệu hóa tài khoản" bị đẩy lệch sang phải, tràn ra ngoài khung nhìn, xuất hiện thanh cuộn ngang. Các menu bên trái chiếm quá nhiều diện tích.', style='List Bullet')
doc.add_paragraph('Mức độ ảnh hưởng: Trung bình (Medium) - Giảm trải nghiệm người dùng (UX) nghiêm trọng trên nền tảng mobile.', style='List Bullet')

doc.add_heading('2. Nguyên nhân (Root Cause)', level=2)
doc.add_paragraph('Lỗi CSS: Trong file `index.css`, cấu trúc Layout (Grid và Flexbox) được set cứng kích thước vật lý thay vì sử dụng tỷ lệ phần trăm hoặc media queries.', style='List Bullet')
doc.add_paragraph('Ví dụ: `.settings-grid` sử dụng `grid-template-columns: 240px 1fr` làm cho cột menu luôn cố định 240px.', style='List Bullet')
doc.add_paragraph('Cách xử lý đã áp dụng: Bổ sung các breakpoint `@media (min-width: ...)` vào file CSS để tự động chuyển đổi cấu trúc grid/flex sang xếp chồng (stacking dọc) trên màn hình mobile.', style='List Bullet')

# Lưu file
out_path = r"C:\Users\DELL\.gemini\antigravity\brain\9fe4ad63-2c63-4528-a2a7-dd02d459bbdd\QA_SRE_Incidents.docx"
doc.save(out_path)
print(f"File saved to {out_path}")
