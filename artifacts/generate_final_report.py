import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Tiêu đề chính
title = doc.add_heading('BÁO CÁO CÔNG VIỆC QA / SRE ENGINEER', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('Dự án: Quản Lý Chi Tiêu (Expense Tracker)')
doc.add_paragraph('Vai trò: QA / SRE Engineer')
doc.add_paragraph('_' * 50)

# ==========================================
# PHẦN 1: TEST HỆ THỐNG
# ==========================================
doc.add_heading('PHẦN 1: NHIỆM VỤ TEST HỆ THỐNG', level=1)
doc.add_paragraph('Tiến hành kiểm thử các luồng chức năng quan trọng của hệ thống dựa trên các test case đã thiết lập.')

def add_tc(doc, title, purpose, precond, steps, expected):
    doc.add_heading(title, level=3)
    doc.add_paragraph(f'Mục đích: {purpose}')
    doc.add_paragraph('Các bước thực hiện:')
    for step in steps:
        doc.add_paragraph(step, style='List Number')
    doc.add_paragraph('Kết quả mong đợi:')
    for exp in expected:
        doc.add_paragraph(exp, style='List Bullet')

add_tc(doc, 'Test Case 1: Đăng nhập hệ thống hợp lệ',
       'Kiểm tra tính năng xác thực người dùng cơ bản.',
       'Tài khoản hợp lệ.',
       ['Truy cập trang Đăng nhập.', 'Nhập email và mật khẩu đúng.', 'Nhấn Đăng nhập.'],
       ['Đăng nhập thành công, chuyển hướng vào Dashboard.'])

add_tc(doc, 'Test Case 2: Kiểm tra tính năng đổi thông tin cá nhân',
       'Kiểm tra khả năng cập nhật Hồ sơ.',
       'Đã đăng nhập.',
       ['Vào Cài đặt -> Tab Hồ sơ.', 'Sửa tên người dùng.', 'Nhấn Lưu.'],
       ['Thông tin được lưu và cập nhật trên UI, không bị lỗi phía server.'])

doc.add_paragraph('_' * 50)

# ==========================================
# PHẦN 2 & 3: DEBUG & XỬ LÝ LỖI + INCIDENT
# ==========================================
doc.add_heading('PHẦN 2 & 3: DEBUG, XỬ LÝ LỖI & 3 INCIDENTS (YÊU CẦU BẮT BUỘC)', level=1)
doc.add_paragraph('Trong quá trình test, debug và cấu hình hệ thống, đã ghi nhận và xử lý 3 sự cố (Incidents) dưới đây:')

# Incident 1
doc.add_heading('Incident 01: Lỗi Crash Backend khi khởi động (ClassNotFoundException)', level=2)
doc.add_heading('Hiện tượng:', level=3)
doc.add_paragraph('Khi thực hiện lệnh `mvn spring-boot:run` để khởi động server backend, tiến trình build thất bại và server bị crash ngay lập tức với mã lỗi `exit code: 1`. Log hệ thống ghi nhận: Caused by: java.lang.ClassNotFoundException: com.example.demo.dto.RegisterRequest.', style='List Bullet')
doc.add_heading('Nguyên nhân:', level=3)
doc.add_paragraph('Thư mục `target` (chứa các file `.class` đã được biên dịch) bị lỗi bộ nhớ đệm (stale cache) khiến Spring Boot không thể tìm thấy file class trên Java Classpath khi chạy. Đã fix bằng cách clean và compile lại.', style='List Bullet')

# Incident 2
doc.add_heading('Incident 02: Xung đột tài nguyên mạng (Port 8081 already in use)', level=2)
doc.add_heading('Hiện tượng:', level=3)
doc.add_paragraph('Tiến trình khởi động Spring Boot thất bại ở giai đoạn khởi tạo Tomcat Web Server. Log hệ thống ghi nhận lỗi: "Web server failed to start. Port 8081 was already in use."', style='List Bullet')
doc.add_heading('Nguyên nhân:', level=3)
doc.add_paragraph('Một tiến trình (Process) ngầm cũ của ứng dụng Spring Boot đã bị crash trước đó nhưng chưa giải phóng tài nguyên mạng, dẫn đến trạng thái zombie "LISTENING" trên cổng 8081. Đã fix bằng lệnh taskkill giải phóng cổng.', style='List Bullet')

# Incident 3
doc.add_heading('Incident 03: Vỡ Layout UI (Overflow) trên thiết bị di động', level=2)
doc.add_heading('Hiện tượng:', level=3)
doc.add_paragraph('Trên trang Cài đặt (Settings), khi người dùng truy cập bằng thiết bị di động, giao diện bị vỡ nghiêm trọng. Khu vực "Vô hiệu hóa tài khoản" bị đẩy lệch sang phải, tràn ra ngoài khung nhìn, xuất hiện thanh cuộn ngang gây lỗi UX.', style='List Bullet')
doc.add_heading('Nguyên nhân:', level=3)
doc.add_paragraph('Lỗi CSS: Trong file `index.css`, cấu trúc Layout (Grid và Flexbox) được set cứng kích thước vật lý (ví dụ grid-template-columns: 240px 1fr) thay vì sử dụng Responsive media queries. Đã fix bằng cách bổ sung @media rules.', style='List Bullet')

# Lưu file
out_path = r"c:\DevOps\QuanLyChiTieu_devops\QuanLyChiTieu_devops\BaoCao_QA_SRE_Full.docx"
doc.save(out_path)
print(f"File saved to {out_path}")
