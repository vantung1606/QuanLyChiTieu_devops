import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Tiêu đề
title = doc.add_heading('Tài Liệu Kịch Bản Kiểm Thử (Test Cases)', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('Dự án: Quản Lý Chi Tiêu (Expense Tracker)', style='List Bullet')
doc.add_paragraph('Vai trò: QA / SRE Engineer', style='List Bullet')
doc.add_paragraph('_' * 50)

def add_tc(doc, title, purpose, precond, steps, expected):
    doc.add_heading(title, level=2)
    doc.add_paragraph(f'Mục đích: {purpose}')
    doc.add_paragraph(f'Tiền điều kiện: {precond}')
    doc.add_paragraph('Các bước thực hiện:')
    for step in steps:
        doc.add_paragraph(step, style='List Number')
    doc.add_paragraph('Kết quả mong đợi:')
    for exp in expected:
        doc.add_paragraph(exp, style='List Bullet')
    doc.add_paragraph('Trạng thái: Chưa test')
    doc.add_paragraph('')

# --- MÔ ĐUN 1 ---
doc.add_heading('MÔ ĐUN 1: XÁC THỰC NGƯỜI DÙNG (AUTHENTICATION)', level=1)

add_tc(doc, 'TC_AUTH_01: Kiểm tra đăng nhập với thông tin hợp lệ',
       'Đảm bảo người dùng có thể đăng nhập thành công khi nhập đúng email/username và mật khẩu.',
       'Tài khoản đã được đăng ký và kích hoạt.',
       ['Truy cập trang Đăng nhập.', 'Nhập Email/Username hợp lệ.', 'Nhập Mật khẩu hợp lệ.', 'Bấm nút "Đăng nhập".'],
       ['Đăng nhập thành công, hệ thống chuyển hướng đến trang Dashboard.', 'Hiển thị thông báo "Đăng nhập thành công".'])

add_tc(doc, 'TC_AUTH_02: Kiểm tra đăng nhập với mật khẩu sai',
       'Đảm bảo hệ thống từ chối đăng nhập khi sai mật khẩu.',
       'Có tài khoản hợp lệ.',
       ['Truy cập trang Đăng nhập.', 'Nhập Email/Username hợp lệ.', 'Nhập sai Mật khẩu.', 'Bấm nút "Đăng nhập".'],
       ['Đăng nhập thất bại, giữ nguyên trang hiện tại.', 'Hiển thị thông báo lỗi "Email hoặc mật khẩu không chính xác".'])

# --- MÔ ĐUN 2 ---
doc.add_heading('MÔ ĐUN 2: CÀI ĐẶT HỒ SƠ & BẢO MẬT (PROFILE & SECURITY)', level=1)

add_tc(doc, 'TC_SET_01: Cập nhật thông tin cá nhân hợp lệ',
       'Kiểm tra tính năng thay đổi họ tên và email.',
       'Đã đăng nhập vào hệ thống.',
       ['Chuyển đến trang Cài đặt -> Tab Hồ sơ.', 'Xóa họ tên cũ, nhập "Nguyễn Văn A".', 'Bấm "Lưu tất cả".'],
       ['Thông tin được cập nhật thành công trên UI.', 'Hiển thị thông báo "Cập nhật hồ sơ thành công!".', 'Reload trang vẫn giữ nguyên thông tin mới.'])

add_tc(doc, 'TC_SET_02: Hiển thị giao diện Responsive trên màn hình di động',
       'Đảm bảo trang Cài đặt không bị vỡ layout trên các thiết bị nhỏ.',
       'Sử dụng trình duyệt thu nhỏ hoặc DevTools (Width < 640px).',
       ['Mở trang Cài đặt.', 'Thu nhỏ màn hình trình duyệt xuống kích thước điện thoại.', 'Kiểm tra hiển thị của khu vực "Vô hiệu hóa tài khoản" và Menu trái.'],
       ['Menu trái chuyển thành thanh cuộn ngang.', 'Khu vực "Vô hiệu hóa tài khoản" tự động xếp chồng (stack) theo chiều dọc.', 'Nút bấm và văn bản hiển thị rõ ràng, không bị tràn.'])

# --- MÔ ĐUN 3 ---
doc.add_heading('MÔ ĐUN 3: XÁC THỰC 2 LỚP (2FA)', level=1)

add_tc(doc, 'TC_2FA_01: Kích hoạt 2FA với mã OTP hợp lệ',
       'Kiểm tra luồng thiết lập bảo mật 2 lớp.',
       'Đã đăng nhập, tính năng 2FA đang tắt.',
       ['Vào Cài đặt -> Tab Bảo mật.', 'Tại mục "Xác thực 2 lớp", bấm "Thiết lập".', 'Hệ thống gửi mã OTP giả lập/thực tế.', 'Nhập mã OTP 6 chữ số hợp lệ vào ô xác nhận.', 'Bấm "Xác nhận".'],
       ['Hiển thị thông báo "Đã bật xác thực 2 lớp thành công!".', 'Trạng thái 2FA trên UI chuyển thành "Đã bật".'])

# Lưu file
out_path = r"c:\DevOps\QuanLyChiTieu_devops\QuanLyChiTieu_devops\Test_Cases_ChiTieu.docx"
doc.save(out_path)
print(f"File saved to {out_path}")
